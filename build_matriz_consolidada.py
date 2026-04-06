"""
build_matriz_consolidada.py
Construye la Matriz Consolidada de Indicadores para todos los programas PARACEL
(vigentes y obsoletos) a partir de los archivos .docx y .xlsx disponibles.
"""

import os
import re
import openpyxl
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side, GradientFill
)
from openpyxl.utils import get_column_letter
from docx import Document
from datetime import datetime

# ── Rutas ────────────────────────────────────────────────────────────────────
BASE = "c:/Users/DiegoMeza/OneDrive - PARACEL S.A/MONITOREO_IMPACTO_SOCIAL_PARACEL/PROGRAMAS_Y_FICHAS_CSS"
DIR_FICHAS   = os.path.join(BASE, "FICHAS _ PROGRAMAS SOCIALES_2026")
DIR_PROGS    = os.path.join(BASE, "PROGRAMAS_FICHAS_CSS")
DIR_OBSOLETOS = os.path.join(DIR_PROGS, "Obsoletos")
OUTPUT       = os.path.join(BASE, "Matriz_Consolidada_TODOS_LOS_PROGRAMAS.xlsx")

# ── Colores PARACEL ───────────────────────────────────────────────────────────
C_VERDE   = "205527"
C_HOJA    = "95BF5B"
C_BLANCO  = "FFFFFF"
C_GRIS    = "F2F2F2"
C_GRIS2   = "E8E8E8"
C_AMARILLO = "FFF2CC"
C_OBSOLETO = "FFE0CC"
C_VERDE_CLARO = "E2EFDA"

# ── Helpers de estilo ─────────────────────────────────────────────────────────
def cell_style(ws, row, col, value="", bold=False, bg=None, fg="000000",
               align="left", wrap=True, size=10, border=True):
    cell = ws.cell(row=row, column=col, value=value)
    cell.font = Font(name="Calibri", bold=bold, size=size, color=fg)
    if bg:
        cell.fill = PatternFill("solid", fgColor=bg)
    cell.alignment = Alignment(
        horizontal=align, vertical="center",
        wrap_text=wrap
    )
    if border:
        thin = Side(style="thin", color="CCCCCC")
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    return cell


def merge_title(ws, row, col_start, col_end, value, bg, fg="FFFFFF", size=12):
    ws.merge_cells(
        start_row=row, start_column=col_start,
        end_row=row, end_column=col_end
    )
    cell = ws.cell(row=row, column=col_start, value=value)
    cell.font = Font(name="Calibri", bold=True, size=size, color=fg)
    cell.fill = PatternFill("solid", fgColor=bg)
    cell.alignment = Alignment(horizontal="center", vertical="center")
    return cell


# ── Parser de nombre de archivo ───────────────────────────────────────────────
def parse_filename(fname):
    """
    Retorna dict con: fecha_raw, anio, codigo, nombre, version, estado
    """
    name = os.path.splitext(fname)[0]
    fecha_raw = ""
    codigo = ""
    version = ""

    # Fecha YYMMDD al inicio
    m = re.match(r'^(\d{6})', name)
    if m:
        fecha_raw = m.group(1)

    # Código PR-XX-XXX
    m2 = re.search(r'(PR-[A-Z]{2}-[A-Z]+)', name, re.IGNORECASE)
    if m2:
        codigo = m2.group(1).upper()

    # Versión v0X
    m3 = re.search(r'v(\d+)', name, re.IGNORECASE)
    if m3:
        version = "v" + m3.group(1).zfill(2)

    # Año
    anio = ""
    if len(fecha_raw) >= 6:
        yy = int(fecha_raw[:2])
        anio = str(2000 + yy)

    # Nombre limpio: quitar fecha, código, versión, guiones extras
    nombre = name
    nombre = re.sub(r'^\d{6}\s*[-–]?\s*', '', nombre)
    nombre = re.sub(r'(PR-[A-Z]{2}-[A-Z]+)\s*[-–]?\s*', '', nombre, flags=re.IGNORECASE)
    nombre = re.sub(r'\s+v\d+\s*$', '', nombre, flags=re.IGNORECASE)
    nombre = re.sub(r'FICHA\s+TÉCNICA\s*[-–]?\s*', '', nombre, flags=re.IGNORECASE)
    nombre = re.sub(r'PROGRAMA\s+DE\s+', 'PROGRAMA DE ', nombre, flags=re.IGNORECASE)
    nombre = nombre.strip(' -–')

    return {
        "fecha_raw": fecha_raw,
        "anio": anio,
        "codigo": codigo,
        "nombre": nombre,
        "version": version,
    }


# ── Extracción desde Word (.docx) ─────────────────────────────────────────────
def extract_docx(fpath):
    """
    Retorna lista de dicts con los datos de indicadores del documento.
    Cada dict: medida, accion, indicadores, medios_verificacion, normas, presupuesto,
               revision_history, objetivos
    """
    doc = Document(fpath)
    rows_out = []
    presupuesto = ""
    objetivos = ""
    responsabilidades = ""
    revision_history = []

    # Objetivos desde párrafos
    capture_obj = False
    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue
        if re.search(r'\bOBJETIVOS?\b', txt, re.IGNORECASE) and len(txt) < 30:
            capture_obj = True
            continue
        if capture_obj and re.search(r'\b(ALCANCE|DEFINICIONES|RESPONSABILIDADES|DESARROLLO)\b', txt, re.IGNORECASE):
            capture_obj = False
        if capture_obj and len(txt) > 10:
            objetivos += txt + " "

    # Tablas
    for table in doc.tables:
        if not table.rows:
            continue

        headers_row = [c.text.strip() for c in table.rows[0].cells]
        headers_str = " ".join(headers_row).lower()

        # Tabla de revisiones históricas
        if "versión" in headers_str or "version" in headers_str:
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and cells[0] and re.match(r'\d+', cells[0]):
                    revision_history.append({
                        "version": cells[0] if cells else "",
                        "fecha": cells[1] if len(cells) > 1 else "",
                        "cambios": cells[2] if len(cells) > 2 else "",
                    })

        # Tabla de presupuesto
        if "presupuesto" in headers_str:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                for c in cells:
                    if re.search(r'USD\s*[\[\d]', c, re.IGNORECASE):
                        presupuesto = c
                        break

        # Tabla principal de indicadores
        if "indicadores" in headers_str and ("medidas" in headers_str or "acciones" in headers_str):
            current_target = ""
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if not any(cells):
                    continue

                # Detectar fila de Target
                first_cell = cells[0] if cells else ""
                if re.match(r'Target', first_cell, re.IGNORECASE) and len(cells) >= 2:
                    # Es encabezado de target
                    current_target = cells[1] if len(cells) > 1 else cells[0]
                    # Si el indicador está en la misma fila (caso donde cells[1] tiene el texto del target y cells[3] tiene indicador)
                    # Verificar si hay indicadores en esta fila también
                    indicadores_cell = cells[3] if len(cells) > 3 else ""
                    medios_cell = cells[4] if len(cells) > 4 else ""
                    normas_cell = cells[5] if len(cells) > 5 else ""
                    if indicadores_cell and indicadores_cell != current_target:
                        rows_out.append({
                            "target": current_target[:200],
                            "medida": cells[1][:200] if len(cells) > 1 else "",
                            "accion": cells[2][:300] if len(cells) > 2 else "",
                            "indicadores": indicadores_cell[:500],
                            "medios_verificacion": medios_cell[:300],
                            "normas": normas_cell[:200],
                        })
                    continue

                # Fila normal de datos
                medida = cells[1] if len(cells) > 1 else ""
                accion = cells[2] if len(cells) > 2 else ""
                indicadores = cells[3] if len(cells) > 3 else ""
                medios = cells[4] if len(cells) > 4 else ""
                normas = cells[5] if len(cells) > 5 else ""

                if indicadores or medida:
                    rows_out.append({
                        "target": current_target[:200],
                        "medida": medida[:200],
                        "accion": accion[:300],
                        "indicadores": indicadores[:500],
                        "medios_verificacion": medios[:300],
                        "normas": normas[:200],
                    })

    return {
        "indicadores": rows_out,
        "presupuesto": presupuesto,
        "objetivos": objetivos.strip()[:500],
        "revision_history": revision_history,
    }


# ── Extracción desde Excel ficha (.xlsx) ─────────────────────────────────────
def extract_ficha_xlsx(fpath):
    """
    Extrae datos de una Ficha Técnica Excel (hojas Propuesta de TC + INDICADORES GENERALES)
    """
    try:
        wb = openpyxl.load_workbook(fpath, data_only=True)
    except Exception as e:
        return {"error": str(e)}

    targets = []
    indicadores_generales = []
    programa_nombre = ""

    # Hoja Propuesta de TC
    if "Propuesta de TC" in wb.sheetnames:
        ws = wb["Propuesta de TC"]
        current_target = {}
        target_num = 0
        for row in ws.iter_rows(values_only=True):
            vals = [str(v).strip() if v is not None else "" for v in row]
            non_empty = [v for v in vals if v]
            if not non_empty:
                continue
            # Nombre del programa (primera fila con texto largo)
            if not programa_nombre and len(non_empty) >= 1:
                candidate = non_empty[0]
                if "FICHA" in candidate.upper() or "PROGRAMA" in candidate.upper():
                    programa_nombre = candidate.replace("FICHA TÉCNICA -", "").replace("FICHA T?CNICA -", "").strip()

            key = vals[1].strip() if len(vals) > 1 else ""
            val = vals[2].strip() if len(vals) > 2 else ""

            if re.match(r'Target\s+\d+', key, re.IGNORECASE):
                if current_target:
                    targets.append(current_target)
                target_num += 1
                current_target = {
                    "target": key,
                    "descripcion": val[:400],
                    "actividades": "",
                    "productos": "",
                    "resultado": "",
                    "presupuesto": "",
                }
            elif key.lower() in ("actividades", "actividades "):
                if current_target:
                    current_target["actividades"] = val[:400]
            elif key.lower() in ("productos",):
                if current_target:
                    current_target["productos"] = val[:400]
            elif key.lower() in ("resultado final", "resultado"):
                if current_target:
                    current_target["resultado"] = val[:300]
            elif key.lower() in ("presupuesto", "prespuesto"):
                if current_target:
                    current_target["presupuesto"] = val[:100]

        if current_target:
            targets.append(current_target)

    # Hoja INDICADORES GENERALES
    if "INDICADORES GENERALES" in wb.sheetnames:
        ws = wb["INDICADORES GENERALES"]
        current_target = ""
        for row in ws.iter_rows(values_only=True):
            vals = [str(v).strip() if v is not None else "" for v in row]
            non_empty = [v for v in vals if v]
            if not non_empty:
                continue
            key = vals[1] if len(vals) > 1 else ""
            val = vals[2] if len(vals) > 2 else ""

            if re.match(r'Target\s+\d+', key, re.IGNORECASE):
                current_target = key
            elif key and len(key) > 5:
                indicadores_generales.append({
                    "target": current_target,
                    "tipo": key[:100],
                    "descripcion": (key + " " + val).strip()[:500],
                })

    return {
        "programa_nombre": programa_nombre,
        "targets": targets,
        "indicadores_generales": indicadores_generales,
    }


# ── Recolección de todos los archivos ────────────────────────────────────────
def collect_all_files():
    """Retorna lista de dicts con metadata de cada archivo encontrado."""
    files = []

    # 1. Fichas Excel 2026
    for fname in sorted(os.listdir(DIR_FICHAS)):
        if not fname.endswith(".xlsx") or fname.startswith("Matriz"):
            continue
        info = parse_filename(fname)
        info["tipo_archivo"] = "FICHA_EXCEL"
        info["estado"] = "VIGENTE"
        info["fpath"] = os.path.join(DIR_FICHAS, fname)
        info["fname"] = fname
        files.append(info)

    # 2. Programas Word vigentes
    for fname in sorted(os.listdir(DIR_PROGS)):
        if not fname.endswith(".docx"):
            continue
        info = parse_filename(fname)
        info["tipo_archivo"] = "PROGRAMA_WORD"
        info["estado"] = "VIGENTE"
        info["fpath"] = os.path.join(DIR_PROGS, fname)
        info["fname"] = fname
        files.append(info)

    # 3. Programas Word obsoletos
    for fname in sorted(os.listdir(DIR_OBSOLETOS)):
        if not fname.endswith(".docx"):
            continue
        info = parse_filename(fname)
        info["tipo_archivo"] = "PROGRAMA_WORD"
        info["estado"] = "OBSOLETO"
        info["fpath"] = os.path.join(DIR_OBSOLETOS, fname)
        info["fname"] = fname
        files.append(info)

    return files


# ══════════════════════════════════════════════════════════════════════════════
# CONSTRUCCIÓN DEL EXCEL
# ══════════════════════════════════════════════════════════════════════════════
def build_excel(files):
    wb = openpyxl.Workbook()

    # ── Hoja 1: Registro de Programas ────────────────────────────────────────
    ws_reg = wb.active
    ws_reg.title = "Registro de Programas"

    # Logo / título
    ws_reg.row_dimensions[1].height = 40
    ws_reg.row_dimensions[2].height = 20
    ws_reg.merge_cells("A1:J1")
    t = ws_reg["A1"]
    t.value = "PARACEL  |  REGISTRO COMPLETO DE PROGRAMAS SOCIALES - TODOS LOS AÑOS"
    t.font = Font("Calibri", bold=True, size=14, color=C_BLANCO)
    t.fill = PatternFill("solid", fgColor=C_VERDE)
    t.alignment = Alignment(horizontal="center", vertical="center")

    ws_reg.merge_cells("A2:J2")
    t2 = ws_reg["A2"]
    t2.value = f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}  |  Vigentes y Obsoletos"
    t2.font = Font("Calibri", size=10, color="666666")
    t2.fill = PatternFill("solid", fgColor=C_GRIS)
    t2.alignment = Alignment(horizontal="center", vertical="center")

    reg_headers = [
        "N°", "AÑO", "CÓDIGO", "VERSIÓN", "NOMBRE DEL PROGRAMA",
        "TIPO ARCHIVO", "ESTADO", "PRESUPUESTO", "OBJETIVOS RESUMIDOS", "ARCHIVO"
    ]
    col_widths = [5, 7, 14, 9, 55, 15, 10, 22, 60, 55]

    for i, h in enumerate(reg_headers, 1):
        cell_style(ws_reg, 3, i, h, bold=True, bg=C_HOJA, fg=C_BLANCO,
                   align="center", size=10)
        ws_reg.column_dimensions[get_column_letter(i)].width = col_widths[i-1]

    ws_reg.freeze_panes = "A4"

    row_n = 4
    for idx, f in enumerate(files, 1):
        estado = f["estado"]
        bg = C_OBSOLETO if estado == "OBSOLETO" else (
            C_VERDE_CLARO if f["tipo_archivo"] == "FICHA_EXCEL" else C_BLANCO
        )

        # Extraer datos básicos
        presupuesto = ""
        objetivos = ""
        if f["tipo_archivo"] == "PROGRAMA_WORD":
            try:
                data = extract_docx(f["fpath"])
                presupuesto = data.get("presupuesto", "")
                objetivos = data.get("objetivos", "")[:200]
            except Exception as e:
                objetivos = f"Error: {e}"
        elif f["tipo_archivo"] == "FICHA_EXCEL":
            try:
                data = extract_ficha_xlsx(f["fpath"])
                presupuestos = [t.get("presupuesto","") for t in data.get("targets",[])]
                presupuesto = " | ".join([p for p in presupuestos if p])
            except Exception as e:
                presupuesto = f"Error: {e}"

        vals = [
            idx, f["anio"], f["codigo"], f["version"], f["nombre"],
            f["tipo_archivo"], estado, presupuesto, objetivos, f["fname"]
        ]
        for ci, v in enumerate(vals, 1):
            bg_use = bg if ci > 1 else C_GRIS2
            cell_style(ws_reg, row_n, ci, v, bg=bg_use,
                       align="center" if ci in (1,2,3,4,6,7) else "left")

        ws_reg.row_dimensions[row_n].height = 30
        row_n += 1

    # ── Hoja 2: Matriz Consolidada ────────────────────────────────────────────
    ws_mat = wb.create_sheet("Matriz Consolidada")

    ws_mat.row_dimensions[1].height = 45
    ws_mat.row_dimensions[2].height = 18
    ws_mat.row_dimensions[3].height = 35

    ws_mat.merge_cells("A1:L1")
    t = ws_mat["A1"]
    t.value = "MATRIZ CONSOLIDADA DE INDICADORES  |  TODOS LOS PROGRAMAS SOCIALES PARACEL"
    t.font = Font("Calibri", bold=True, size=14, color=C_BLANCO)
    t.fill = PatternFill("solid", fgColor=C_VERDE)
    t.alignment = Alignment(horizontal="center", vertical="center")

    ws_mat.merge_cells("A2:L2")
    t2 = ws_mat["A2"]
    t2.value = "Vigentes y Obsoletos  |  Extraído de documentos FICHAS y PROGRAMAS"
    t2.font = Font("Calibri", size=9, color="666666")
    t2.fill = PatternFill("solid", fgColor=C_GRIS)
    t2.alignment = Alignment(horizontal="center", vertical="center")

    mat_headers = [
        "N°", "AÑO", "CÓDIGO", "VERSIÓN", "ESTADO",
        "PROGRAMA", "TARGET / MEDIDA PROPUESTA",
        "ACCIONES CONCRETAS", "INDICADORES",
        "MEDIOS DE VERIFICACIÓN", "NORMAS VINCULADAS",
        "POBLACIÓN ASOCIADA"
    ]
    mat_widths = [5, 7, 14, 9, 10, 45, 45, 45, 50, 40, 35, 30]

    for i, h in enumerate(mat_headers, 1):
        cell_style(ws_mat, 3, i, h, bold=True, bg=C_VERDE, fg=C_BLANCO,
                   align="center", size=10)
        ws_mat.column_dimensions[get_column_letter(i)].width = mat_widths[i-1]

    ws_mat.freeze_panes = "A4"

    row_n = 4
    indicator_count = 0

    for f in files:
        if f["tipo_archivo"] != "PROGRAMA_WORD":
            continue  # Indicadores provienen de Word docs
        try:
            data = extract_docx(f["fpath"])
        except Exception as e:
            print(f"  ERROR {f['fname']}: {e}")
            continue

        estado = f["estado"]
        bg = C_OBSOLETO if estado == "OBSOLETO" else C_VERDE_CLARO

        indicadores = data.get("indicadores", [])
        if not indicadores:
            # Fila vacía para documentar que el programa existe
            indicadores = [{"target": "", "medida": "", "accion": "",
                            "indicadores": "(sin indicadores estructurados)",
                            "medios_verificacion": "", "normas": ""}]

        for ind in indicadores:
            indicator_count += 1
            vals = [
                indicator_count,
                f["anio"],
                f["codigo"],
                f["version"],
                estado,
                f["nombre"],
                ind.get("target", "") or ind.get("medida", ""),
                ind.get("accion", ""),
                ind.get("indicadores", ""),
                ind.get("medios_verificacion", ""),
                ind.get("normas", ""),
                ind.get("target", ""),  # población/target
            ]
            for ci, v in enumerate(vals, 1):
                cell_style(ws_mat, row_n, ci, v,
                           bg=bg if ci > 1 else C_GRIS2,
                           align="center" if ci in (1,2,3,4,5) else "left",
                           size=9)
            ws_mat.row_dimensions[row_n].height = 45
            row_n += 1

    # ── Hoja 3: Fichas Técnicas 2026 ─────────────────────────────────────────
    ws_fic = wb.create_sheet("Fichas Técnicas 2026")

    ws_fic.row_dimensions[1].height = 40
    ws_fic.merge_cells("A1:K1")
    t = ws_fic["A1"]
    t.value = "FICHAS TÉCNICAS 2026  |  PROPUESTA DE TEORÍA DE CAMBIO POR PROGRAMA"
    t.font = Font("Calibri", bold=True, size=13, color=C_BLANCO)
    t.fill = PatternFill("solid", fgColor=C_HOJA)
    t.alignment = Alignment(horizontal="center", vertical="center")

    fic_headers = [
        "N°", "CÓDIGO", "PROGRAMA", "TARGET N°", "DESCRIPCIÓN TARGET",
        "ACTIVIDADES CLAVE", "PRODUCTOS", "RESULTADO FINAL",
        "PRESUPUESTO", "AÑO", "VERSIÓN"
    ]
    fic_widths = [5, 14, 45, 9, 55, 55, 45, 45, 18, 7, 9]

    for i, h in enumerate(fic_headers, 1):
        cell_style(ws_fic, 2, i, h, bold=True, bg=C_HOJA, fg=C_BLANCO,
                   align="center", size=10)
        ws_fic.column_dimensions[get_column_letter(i)].width = fic_widths[i-1]

    ws_fic.freeze_panes = "A3"

    row_n = 3
    fic_count = 0
    for f in files:
        if f["tipo_archivo"] != "FICHA_EXCEL":
            continue
        try:
            data = extract_ficha_xlsx(f["fpath"])
        except Exception as e:
            print(f"  ERROR ficha {f['fname']}: {e}")
            continue

        targets = data.get("targets", [])
        if not targets:
            targets = [{"target": "", "descripcion": "(sin targets)",
                        "actividades": "", "productos": "", "resultado": "",
                        "presupuesto": ""}]

        bg_alt = C_GRIS if fic_count % 2 == 0 else C_BLANCO

        for t_idx, targ in enumerate(targets):
            fic_count += 1
            t_num = re.search(r'\d+', targ.get("target",""))
            t_num_str = t_num.group(0) if t_num else str(t_idx+1)
            vals = [
                fic_count,
                f["codigo"],
                f["nombre"],
                t_num_str,
                targ.get("descripcion",""),
                targ.get("actividades",""),
                targ.get("productos",""),
                targ.get("resultado",""),
                targ.get("presupuesto",""),
                f["anio"],
                f["version"],
            ]
            for ci, v in enumerate(vals, 1):
                cell_style(ws_fic, row_n, ci, v,
                           bg=bg_alt if ci > 1 else C_GRIS2,
                           align="center" if ci in (1,2,4,10,11) else "left",
                           size=9)
            ws_fic.row_dimensions[row_n].height = 60
            row_n += 1

    # ── Hoja 4: Resumen por Programa ─────────────────────────────────────────
    ws_res = wb.create_sheet("Resumen por Programa")

    ws_res.row_dimensions[1].height = 40
    ws_res.merge_cells("A1:I1")
    t = ws_res["A1"]
    t.value = "RESUMEN EJECUTIVO  |  TODOS LOS PROGRAMAS PARACEL (VIGENTES Y OBSOLETOS)"
    t.font = Font("Calibri", bold=True, size=13, color=C_BLANCO)
    t.fill = PatternFill("solid", fgColor=C_VERDE)
    t.alignment = Alignment(horizontal="center", vertical="center")

    res_headers = [
        "CÓDIGO", "PROGRAMA", "VERSIONES DISPONIBLES",
        "AÑOS DOCUMENTADOS", "ESTADO ACTUAL",
        "N° TARGETS (ÚLTIMA VERSIÓN)", "PRESUPUESTO (ÚLTIMA VERSIÓN)",
        "TIPO DE PROGRAMA", "OBSERVACIONES"
    ]
    res_widths = [14, 55, 30, 25, 14, 20, 25, 18, 35]

    for i, h in enumerate(res_headers, 1):
        cell_style(ws_res, 2, i, h, bold=True, bg=C_VERDE, fg=C_BLANCO,
                   align="center", size=10)
        ws_res.column_dimensions[get_column_letter(i)].width = res_widths[i-1]

    ws_res.freeze_panes = "A3"

    # Agrupar por código
    from collections import defaultdict
    by_code = defaultdict(list)
    for f in files:
        key = f["codigo"] if f["codigo"] else f["nombre"][:30]
        by_code[key].append(f)

    row_n = 3
    for code, flist in sorted(by_code.items()):
        # Determinar estado actual
        vigentes = [f for f in flist if f["estado"] == "VIGENTE"]
        estado_actual = "VIGENTE" if vigentes else "OBSOLETO"

        # Versiones y años
        versions = sorted(set(f["version"] for f in flist if f["version"]))
        anios = sorted(set(f["anio"] for f in flist if f["anio"]))

        # Datos de última versión (mayor número de versión)
        latest = sorted(flist, key=lambda x: x["version"])[-1]
        nombre = latest["nombre"]

        # Targets y presupuesto de última versión Word
        n_targets = ""
        presupuesto = ""
        word_latest = [f for f in flist if f["tipo_archivo"]=="PROGRAMA_WORD"]
        if word_latest:
            wl = sorted(word_latest, key=lambda x: x["version"])[-1]
            try:
                data = extract_docx(wl["fpath"])
                presupuesto = data.get("presupuesto","")
                # Contar targets únicos en indicadores
                targets_found = set()
                for ind in data.get("indicadores",[]):
                    tgt = ind.get("target","")
                    if tgt:
                        targets_found.add(tgt[:30])
                n_targets = str(len(targets_found)) if targets_found else ""
            except:
                pass

        # Tipo de programa
        tipo = "SOCIAL" if "SS" in code else ("TALENTO HUMANO" if "TH" in code else "GESTIÓN")

        bg = C_OBSOLETO if estado_actual == "OBSOLETO" else C_VERDE_CLARO

        vals = [
            code, nombre,
            ", ".join(versions),
            ", ".join(anios),
            estado_actual,
            n_targets,
            presupuesto,
            tipo,
            ""
        ]
        for ci, v in enumerate(vals, 1):
            cell_style(ws_res, row_n, ci, v,
                       bg=bg,
                       bold=(estado_actual == "VIGENTE" and ci == 2),
                       align="center" if ci in (1,3,4,5,6,8) else "left",
                       size=10)
        ws_res.row_dimensions[row_n].height = 35
        row_n += 1

    # ── Guardar ───────────────────────────────────────────────────────────────
    wb.save(OUTPUT)
    print(f"\n✓ Archivo guardado: {OUTPUT}")
    print(f"  Hojas: {[s.title for s in wb.worksheets]}")


# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Recolectando archivos...")
    files = collect_all_files()
    print(f"  Total archivos encontrados: {len(files)}")
    for f in files:
        print(f"  [{f['estado'][:3]}] {f['anio']} {f['codigo']} {f['version']} - {f['fname'][:60]}")

    print("\nConstruyendo Excel...")
    build_excel(files)
